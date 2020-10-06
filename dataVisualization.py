#This file generates a word document with information summarizing the data put into it

#We will begin with loading and cleaning/updating the dataframe.
# Afterwards, we will begin visualizing the relevant data.
import datetime
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import os
import derivitizationUtils as du #Has functions such as basic analysis/load dataframe
from itertools import combinations

#Libraries necessary to generate the MS Word Document
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH #Used to align str(number) in cells

#Import the graph package
plt.rcParams.update({'figure.max_open_warning':0})

#Create a directory for figures/plots if it doesn't already exist
if not os.path.exists('./figures'):
    os.makedirs('./figures')

start_time = datetime.datetime.now()

#Load the dataframe
data_folder = 'D:/J4LMData/Aggregates/'
data_filename = 'J4LM_Aggregate_AFRICOM'
#data_extension = '.txt'
data_extension = '.csv'
encoding = 'latin1' #File encoding type (examples: 'latin1', 'utf_16', etc.).  Will automatically be determined if not specified


#Determine file encoding
if encoding == None:
    chardet_encoding = du.determine_encoding(data_folder, data_filename, data_extension)
    encoding = du.encoding_string(chardet_encoding)


#Read an excel file
if data_extension == '.xlsx':
    try:
        df = pd.read_excel(data_folder + data_filename + data_extension)
        raw_data_file = data_folder + data_filename + data_extension
        print('File loaded.')
    except:
        print('Error: Data file not found.')
    print("Load Data Run Time: ", datetime.datetime.now() - start_time)
    
#Read other file types (.csv, .txt)
else:
    try:
        df = du.load_dataframe(data_folder, data_filename, data_extension, encode_type=encoding, separator='|')
        raw_data_file = data_folder + data_filename + data_extension
        print('File loaded.')
    except:
        print('Error: Data file not found.')
    print("Load Data Run Time: ", datetime.datetime.now() - start_time)
print(df.columns)

#Options set to improve display
pd.options.display.float_format = '{:.2f}'.format #Suppressing scientific notation
pd.options.mode.chained_assignment = None #Disable some warnings from pandas
pd.set_option('display.max_columns', 250) #Set max number of columns to display in notebook

print("Load Data Run Time: ", datetime.datetime.now() - start_time)



################################################################
#ESTABLISH FUNCTION
def report_generator(df, data_elements = None, title = True, subtitle = ['date', 'categorical', 'measure'], corrplots = True, binaryplots = False):
    #df = dataframe from above that we are analyzing
    #data_elements = list of lists defining which df columns are which type of data (OPTIONAL)
        #NOTE: First list is always dates, second list is always categorical vars, third list is always measure vars
        #NOTE: Always 3 lists, even if one is empty
        # ex: [['DATE_COL1', 'DATE_COL2'], ['CAT_COL1'], []]
    #title = if present, creates a distint title page + introduction page (pages 1+2) (OPTIONAL)
        # default = True, this will create title/intro pages.  Set to false to skip the title pages.
    #subtitle = list containing combination of 'date', 'categorical', and 'measure' strings (OPTIONAL)
        # ex: ['date', 'measure']  ex2: ['categorical']
        # The above first example would then only print out individual summaries for date and measure columns
        # default is to print summaries for all three categories of data
    #corrplots = if True, will generate plots of all numerical columns that appear to be correlated (values > 0.5 or < -0.5)
    #binaryplots = if True, will generate plots comparing categorical columns that have 15 or less unique values across the dataset
        # Note: For improved runtimes, only samples of the columns are taken for the plotting, rather than the entire columns
        
    start_time = datetime.datetime.now()
    
    #Get number of columns in input dataframe
    input_cols = df.shape[1]
    
    #Basic verification that inputs for the function are good
    if data_elements != None:
        #Check that submitted list has 3 lists inside
        if len(data_elements) != 3:
            print('Please submit data_elements as a list of 3 lists.')
            return None  
        #Check that submitted column variables are all present in the dataframe
        dataframecolumnlist = list(df.columns.values)
        for innerlist in data_elements: #Loop over 3 inner lists
            for entry in innerlist: #Loop over entries in inner list
                if type(entry) != str:
                    print('Only submit strings for column names in data_elements. ' + str(entry) + ' is not a string')
                    return None
                if entry not in dataframecolumnlist:
                    print('Submitted column names in data_elements not present in dataframe: ' + str(entry))
                    return None       
    
    #Check that title == True or title == False
    if title != True and type(title) != str:
        print('Please input either True or a string value for the title varible in report_generator')
        return None
    
    #Check that subtitle only contains a list with the 3 values 'date', 'categorical', or 'measure'
    subtitle_list = ['date', 'categorical', 'measure']
    if type(subtitle) is not list:
        print('Please use a list to define the subtitble variable in report_generator. '
             ' Ex: ["date"], Ex2: ["categorical", "measure"] ')
        return None
    if (set(subtitle) <= set(subtitle_list)) == False:
        print('The subtitle variable in report_generator can only contain the following values: ' 
             'date, categorical, measure')
        return None
    
    

    ################################################################
    #DATA CLEANING AND AUTOMATIC DATA TYPE CATEGORIZATION
    
    #Ensure that column names do not contain special characters to prevent bugs with histogram creation
    # Note: bugs arise from the fact that files cannot be saved with these characters in their names
    for column in df:
        if '/' in str(column):
            print("Renaming column " + str(column) + " to replace slashes (/) with underscores (_).")
        if '|' in str(column):
            print("Renaming column " + str(column) + " to replace pipes (|) with underscores (_).")
        if '?' in str(column):
            print("Renaming column " + str(column) + " to replace question marks (?) with underscores (_).")
        if '*' in str(column):
            print("Renaming column " + str(column) + " to replace asterisks (*) with underscores (_).")
        if ':' in str(column):
            print("Renaming column " + str(column) + " to replace colons (:) with underscores (_).")
        if '>' in str(column):
            print("Renaming column " + str(column) + " to replace greater thans (>) with underscores (_).")
        if '<' in str(column):
            print("Renaming column " + str(column) + " to replace less thans (<) with hyphens (-).")
    df.columns = df.columns.str.replace("/", "_")
    df.columns = df.columns.str.replace("|", "_")
    df.columns = df.columns.str.replace("?", "_")
    df.columns = df.columns.str.replace("*", "_")
    df.columns = df.columns.str.replace(":", "_")
    df.columns = df.columns.str.replace(">", "_")
    df.columns = df.columns.str.replace("<", "-")

    #If data types were not submitted manually: For each column, determine data format, then add to appropriate list
    if data_elements == None:
        #Establish empty lists
        measure_fields = []
        date_fields = []
        categorical_fields = []
        nan_fields = []

        #Loop over columns, add columns to correct list (measure/date/categorical fields)
        rowcheck_threshold = 2000 #Set number of rows to investigate when determining column data type ##
        for column in df:
            #Check if column is entirely filled with NaN's.  If so, no point cleaning
            if pd.isnull(df[column]).sum() == len(df[column]):
                print('Column is entirely NaN, excluding from analysis: ' + str(column))
                nan_fields.append(column)
            #Check if column is numerical ('measure' field)
            if (np.issubdtype(df[column].dtype, np.number) == True) and (column not in nan_fields):
                measure_fields.append(column)
                print('Column is numerical: ' + str(column))
            #Check if column is bool; if so treat as categorical
            if df[column].dtype == 'bool':
                categorical_fields.append(column)
                print('Column is bools (True/False), treated as categorical: ' + str(column))

            #Check if column is numerical but stored as strings (or has a few string typos that coverted column type to object)
            if column not in nan_fields and column not in measure_fields and column not in categorical_fields:
                numsuccess = 0 #Used to track number of successful numerical checks
                numfail = 0 #Used to track number of unsuccessful numerical checks
                for index, row in df.head(rowcheck_threshold).iterrows():
                    #We will check first n rows (set by rowcheck_threshold) and make sure they are 99% consistent with numerical values
                    if pd.isnull(df[column][index]) == False:
                        try:
                            df[column][index].isdigit() == True #Try to use .isdigit() ('1' == True, 1 -> error, 'One' -> False)
                        except:
                            numsuccess = numsuccess + 1  #.isdigit() will fail for ints/floats
                        else:
                            if df[column][index].isdigit() == True:
                                numsuccess = numsuccess + 1 
                            else:
                                numfail = numfail + 1
                #Check to see that 99%+ of the column values return numerical values        
                if numsuccess != 0 or numfail != 0: #Avoid division by 0
                    if (numsuccess/(numsuccess + numfail)) >= 0.99:
                        measure_fields.append(column)
                        print('Column is numerical: ' + str(column))

            #Now determine if column is datetime or categorical
            if column not in nan_fields and column not in measure_fields and column not in categorical_fields:
                datesuccess = 0 #Used to track number of successful datetime conversions
                datefail = 0 #Used to track number of unsuccessful datetime conversions
                for index, row in df.head(rowcheck_threshold).iterrows():
                    #We will check first n rows (set by rowcheck_threshold) and make sure they are 90% consistent with datetime standards
                    if pd.isnull(df[column][index]) == False: #Only test non-null values
                        try:
                            pd.to_datetime(df[column][index]) #See if value can be converted to datetime
                        except:
                            datefail = datefail + 1 #If conversion fails, increment datefail
                        else:
                            #Make sure date makes sense; we shouldn't have pre-1990 data
                            if pd.to_datetime(df[column][index]) > pd.to_datetime('1990-01-01'): 
                                datesuccess = datesuccess + 1
                            else:
                                datefail = datefail + 1 #If date is pre 1990, doesn't make sense (ex: '1' -> '1970-01-01')
                #Check to see that 90%+ of the column values return good datetimes
                if datesuccess != 0 or datefail != 0: #Avoid division by 0; this only happens if first n entries (set by rowcheck_threshold) are NaN
                    if (datesuccess/(datesuccess + datefail)) >= 0.9:
                        date_fields.append(column)
                        print('Column is datetime: ' + str(column))
                    else:
                        categorical_fields.append(column)
                        print('Column is categorical: ' + str(column))
                else: #If first n rows are all NaN (datesuccess and datefail == 0), exclude column from analysis
                    print('All tested entries are NaN, excluding column from analysis: ' + str(column))
                    nan_fields.append(column)
                    del df[column]
        
        #Only analyze columns that are in categories in subtitle field (measure, date or category)
        for column in df:
            if (column in date_fields and 'date' not in subtitle) or (column in categorical_fields and 'categorical' not in subtitle) or (column in measure_fields and 'measure' not in subtitle):
                print('Excluding column from analysis: ' + str(column))
                del df[column]
            elif column in nan_fields: #Exclude fully NaN columns
                #print('Excluding fully NaN columns from analysis: ' + str(column)) ##These columns already printed when assigned to nan_fields
                del df[column]        
                
        #If column is identified as measure, but has <30 unique values, turn it into a category
        for column in measure_fields:
            if len(df[column].unique()) <= 30:
                print('Column consistent with measure column, but coverted to category due to low (<30) unique values: ', column)
                measure_fields.remove(column)
                categorical_fields.append(column)
                
    
    else: #If data column types were already determined/submitted, simply assign them to correct list
        date_fields = data_elements[0]
        categorical_fields = data_elements[1]
        measure_fields = data_elements[2]
        nan_fields = []
        
        #Remove slashes in column names if they're present
        for index, field in enumerate(date_fields):
            date_fields[index] = field.replace("/", "_")
        for index, field in enumerate(categorical_fields):
            categorical_fields[index] = field.replace("/", "_")
        for index, field in enumerate(measure_fields):
            measure_fields[index] = field.replace("/", "_")
                
        # Only analyze columns manually submitted about dataframe by removing columns that weren't submitted
        for column in df:
            if (column not in date_fields) and (column not in categorical_fields) and (column not in measure_fields):
                print('Excluding column from analysis: ' + str(column))
                del df[column]
                MemUsage = df.memory_usage().sum() #Keeps memory usage down
            elif pd.isnull(df[column]).sum() == len(df[column]):
                print('Column is entirely NaN, excluding from analysis: ' + str(column))
                nan_fields.append(column)
                for entry in [date_fields, categorical_fields, measure_fields]:
                    if column in entry:
                        entry.remove(column)
                del df[column]
                MemUsage = df.memory_usage().sum() #Keeps memory usage down

        
    #Clean columns in each list
    #Date columns must always be cleaned to convert to datetime formatting
    print("Cleaning date data")
    for field in date_fields:
        du.clean_date_data(df, [field])
        
    cleanfile = True #True by default, only switch if category/measure fields have already been cleaned
    if cleanfile == True: #If you want to clean the measure/category fields
        print("Cleaning measure data")
        for field in measure_fields:
            du.clean_measure_data(df,[field])
        print("Cleaning categorical data")
        for field in categorical_fields:
            du.clean_categorical_data(df, [field])
            
    #If cleaning caused a column to become fully null, exclude it
    for column in df:
        #Check if column is entirely filled with NaN's.  If so, no point analyzing
        if pd.isnull(df[column]).sum() == len(df[column]):
            print('After cleaning, column is entirely NaN, excluding from analysis: ' + str(column))
            nan_fields.append(column)
            if column in date_fields:
                date_fields.remove(column)
            elif column in categorical_fields:
                categorical_fields.remove(column)
            elif column in measure_fields:
                measure_fields.remove(column)
            del df[column]
    
    #Option to save the file (in order to avoid cleaning in future runs)
    savefile = False
    if savefile == True:
        du.save_file(df, data_folder, data_filename, False)

    end_time = datetime.datetime.now()
    print("Run Time - ", end="", flush=True)
    if data_elements == None:
        print("Column Type Assignment & ", end="", flush=True)
    if savefile == True:
        print("File Saving & ", end="", flush=True)
    print("Data Cleaning: ", end="", flush=True)
    print(end_time-start_time)
    start_time = datetime.datetime.now()



    #################################################################
    #INITIAL ANALYSES
    #Creation of a dataframe (data_qlt_df) that contains summarized information about each column

    #Number of rows
    no_of_rows = len(df.columns)

    #Constructing the data_qlt_df dataframe and pre-assigning and columns.
    # Pre-assigning the number of rows the dataframe would have is memory and processing efficient.
    # This is a better approach than continuous append or concat operation to dataframe
    data_qlt_df = pd.DataFrame(index=np.arange(0, no_of_rows),
                                 columns=('column_name', 'col_data_type','non_null_values',
                                          'unique_values_count', 'column_dtype'))

    #Add rows to the data_qlt_df dataframe
    for ind, cols in enumerate(df.columns):
        # Count of unique values in the column
        col_unique_count = df[cols].nunique()
        data_qlt_df.loc[ind] = [cols,
                                df[cols].dtype,
                                df[cols].count(),
                                col_unique_count,
                                cols + '~'+ str(df[cols].dtype)]

    #Use describe() to get column stats of raw dataframe. This will be merged with the DPD.
    raw_num_df = df.describe().T.round(2)


    #Key step:
    # Merge the df.describe() output with rest of the info to create a single Data Profile Dataframe
    data_qlt_df = pd.merge(data_qlt_df, raw_num_df, how='left', left_on='column_name', right_index=True)

    #Add more columns to our summary dataframe
    #Calculate percentage of non-null values over total number of values
    data_qlt_df['%_of_non_nulls'] = (data_qlt_df['non_null_values']/df.shape[0])*100
    #Calculate null values for the column
    data_qlt_df['null_values'] = df.shape[0] - data_qlt_df['non_null_values']
    #Calculate percentage of null values over total number of values
    data_qlt_df['%_of_nulls'] = 100 - data_qlt_df['%_of_non_nulls']
    #Calculate the total count of column values
    data_qlt_df["count"] = data_qlt_df['null_values'] + data_qlt_df['non_null_values']
    #Column that has the dtype, which now includes 'datetime' as an option
    data_qlt_df["data_type"] = data_qlt_df['col_data_type']
    for index, row in data_qlt_df.iterrows():
        #If the column is actually a datetime, this value will be updated from 'object' to 'datetime'
        if data_qlt_df['column_name'][index] in date_fields:
            data_qlt_df['data_type'][index] = 'object - datetime'
    #Column that is labeled 'measure', 'date', or 'categorical', for which type the data falls into 
    data_qlt_df["input_type"] = np.nan
    for index, row in data_qlt_df.iterrows():
        if data_qlt_df['column_name'][index] in date_fields:
            data_qlt_df['input_type'][index] = 'date'
        if data_qlt_df['column_name'][index] in measure_fields:
            data_qlt_df['input_type'][index] = 'measure'
        if data_qlt_df['column_name'][index] in categorical_fields:
            data_qlt_df['input_type'][index] = 'categorical'

    #Reorder the Data Profile Dataframe columns while adding min/max and other useful columns
    #Note: 2 cases: Having numerical columns and not having numerical columns
    # If numerical columns already exist, previous merged df will have min, max, 25%, 50%, etc. columns already existing
    if measure_fields == []: #No numerical columns in input dataframe
        data_qlt_df = data_qlt_df[['column_name', 'col_data_type', 'non_null_values', '%_of_non_nulls', 'null_values', '%_of_nulls', 'unique_values_count', 'count', 'data_type', 'input_type']]
        data_qlt_df['min'] = np.nan #Initialize new columns
        data_qlt_df['max'] = np.nan
        data_qlt_df['mean'] = np.nan
        for column in df: #Loop over columns in df
            for index, row in data_qlt_df.iterrows(): #Loop over indicies in our new df
                #Update the min/max values for date entries
                currentcol = data_qlt_df['column_name'][index]
                if data_qlt_df['column_name'][index] in date_fields: #Check if column is a date field
                    data_qlt_df['min'][index] = (df[currentcol]).min()
                    data_qlt_df['max'][index] = (df[currentcol]).max() 
                    #Implement date mean with the following function:
                    data_qlt_df['mean'][index] = ((df[currentcol]) - (df[currentcol]).min()).mean() + (df[currentcol]).min()
                #Update the min/max values for category entries
                if data_qlt_df['column_name'][index] in categorical_fields: #Check if column is a category field
                    #Min and max values for categories are the min/max for their length as strings (ex: 'UN1' -> 3)
                    data_qlt_df['min'][index] = df[currentcol].str.len().min()
                    data_qlt_df['max'][index] = df[currentcol].str.len().max()
                    data_qlt_df['mean'][index] = df[currentcol].str.len().mean()
        
    elif measure_fields != []: #If measure fields are present, 25%, 50%, etc. columns will exist
        data_qlt_df = data_qlt_df[['column_name', 'col_data_type', 'non_null_values', '%_of_non_nulls', 'null_values', '%_of_nulls', 'unique_values_count', 'count', 'mean', 'std', 'min', '25%', '50%', '75%', 'max', 'data_type', 'input_type']]
        for index, row in data_qlt_df.iterrows(): #Loop over indicies in our new df
            #Update the min/max values for date entries
            currentcol = data_qlt_df['column_name'][index]
            if currentcol in date_fields: #Check if column is a date field
                data_qlt_df['min'][index] = (df[currentcol]).min()
                data_qlt_df['max'][index] = (df[currentcol]).max()
                #Implement date mean with the following function:
                data_qlt_df['mean'][index] = ((df[currentcol]) - (df[currentcol]).min()).mean() + (df[currentcol]).min()
            #Update the min/max values for category entries
            if currentcol in categorical_fields: #Check if column is a category field
                #Min and max values for categories are the min/max for their length as strings (ex: 'UN1' -> 3)
                data_qlt_df['min'][index] = df[currentcol].str.len().min()
                data_qlt_df['max'][index] = df[currentcol].str.len().max()
                data_qlt_df['mean'][index] = df[currentcol].str.len().mean()                                 
    
    end_time = datetime.datetime.now()
    print("Run Time - Initial Analyses: ", end_time - start_time)
    start_time = datetime.datetime.now()
    


    ##################################################
    #CONTINUED ANALYSES
    #Creation of histograms for each column and creation of null_vals_df for tracking columns with high null counts
    
    #Get the list of date/time columns from raw dataframe that are not empty, create histograms for these
    if date_fields != [] and 'date' in subtitle:
        du.histogram_analysis(df, 'figures', date_fields, isdate = True)
        
    #Get the list of object columns from raw dataframe that are not empty, create histograms for these
    if categorical_fields != [] and 'categorical' in subtitle:
        du.histogram_analysis(df, 'figures', categorical_fields)

    #Get the list of numeric columns from raw dataframe that are not empty, create histograms for these
    if measure_fields != [] and 'measure' in subtitle: #If there's at least one measure field
        du.numeric_column_plot(df, 'figures', measure_fields)
        df.drop(u'quartiles',axis=1,inplace=True) #Deletes an unnecessary column from the df
        
    #Create correlation plots for correlated numerical data
    if corrplots == True:
        measure_pair_list = []
        for combo in combinations(measure_fields, 2): #Look at all combinations of two measure values
            
            #Create a list with the unique combination of two fields
            measure_pair = []
            for value in combo:
                measure_pair.append(value)
            
            #Find correlation value of the two columns, append combo if correlation is larger than 0.5
            corrtable = df[measure_pair].corr(method = 'pearson')
            corrvalue = corrtable[measure_pair[0]][1]
            if (corrvalue >= 0.5 or corrvalue <= -0.5) and np.isnan(corrvalue) == False: 
                #Append combo to list of lists containing each correlated pair
                measure_pair_list.append(measure_pair)
        #Create plots for each combo in list of correlated columns
        print(str(len(measure_pair_list)) + " highly correlated/anticorrelated measure pairs found.")
        du.corrplot(df, 'figures', measure_pair_list)
        
    #Create internal binary plots for categorical values with 15 or less unique values in the column 
    if binaryplots == True:
        #Assign variables for plotting function
        unique_cutoff = 100
        sample_size = 10000
        #Create list of categorical columns with 15 or less unique values
        binarycatlist = []
        for field in categorical_fields:
            unique_count = data_qlt_df.at[data_qlt_df['column_name'].eq(field).idxmax(), 'unique_values_count']
            if unique_count <= 15: #If 15 or less unique values in the column
                binarycatlist.append(field) #Add this field to list of categories to plot
        
        #Create plots.  unique_cutoff -> don't make more than x plots; sample_size -> number of rows to sample from for plot
        if len(binarycatlist) >= 2: #Only create plots if 2 or more categories to compare
            #Print how many plots we are generating
            plottracker = 0
            for i in range(0,len(binarycatlist)):  plottracker = plottracker+i
            print("Generating " + str(plottracker) + " unique categorical binary plots")
            #Create binary categorical plots
            du.internal_category_binary_analysis(df, 'figures', binarycatlist, unique_cutoff, sample_size)
        else:
            print("Insufficient categorical columns with low enough unique counts for binary plotting")

    #Set empty values threshold.  Can be adjusted later for differently-sized datasets 
    threshold_perc = 0.75
    col_vals_threshold = df.shape[0] * threshold_perc #Number of rows * threshold_percentage

    #Create a df to track null values 
    null_vals_df = data_qlt_df[data_qlt_df['non_null_values'] < col_vals_threshold][['column_name', 'data_type', 'unique_values_count', 'min', 'max', 'non_null_values', 'null_values', '%_of_nulls']]

    end_time = datetime.datetime.now()
    print("Run Time - Continued Analyses: ", end_time - start_time)
    start_time = datetime.datetime.now()



    ###################################################################
    #BEGIN DOCUMENT GENERATION

    #Create Document object
    document = Document()
    
    #Set margins to 0.75
    sections = document.sections
    for section in sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    #Page 1
    #Add Title
    if title == True: #Only generate the title and intro pages if function input (title) = True
        date_generate = datetime.datetime.now().date() #Only show date, not time
        document.add_heading('Data Profile Dataframe: ' + str(date_generate), 0)
        document.add_heading(raw_data_file, 0) #Show name of raw datafile

        #Cover page paragraph
        p = document.add_paragraph('The main objective of this notebook is ')
        p.add_run('only').bold = True
        p.add_run(' to understand the raw data profile.  For example, we would like to explore data types, '
            'minimum and maximum values, ranges, unique values, etc.')
        p = document.add_paragraph('This notebook does not make dataset-specific suggestions such as how to clean the data '
            'and what data transformations should be performed based on the data profile.')
        p = document.add_paragraph('')
        p.add_run('This notebook is kept generic so that it can be used with a wide range of datasets.').italic = True
        document.add_page_break()


        #Page 2
        #Heading 1
        document.add_heading('The Game Changer - Data Profile Dataframe (DPD)', level=1)
        p = document.add_paragraph('The game changer for exploratory data analysis is the final')
        p.add_run(' Data Profile Dataframe').bold = True
        p.add_run(' that is generated which combines ')
        p.add_run('all').bold = True
        p.add_run(' the information required to inform data cleansing and optimisation (memory and processing) '
                'decisions. Instead of using various Pandas commands at different instances and going back and forth to cross '
                'refer information, a Data Profile Dataframe brings all information into a single dataframe. This will be '
                'useful when reviewing the data profile with the business subject matter or other team members as all '
                'information related to data profile is in a single, easy-to-understand format.')

        p = document.add_paragraph('Understanding the data is ')
        p.add_run('the critical step').bold = True
        p.add_run(' in preparing the data to be used for analytics. As many experts will point out, '
                'the data preparation and transformation into a tidy format often requires ~80% of the effort '
                'in any data analytics or data analysis project.')
        p = document.add_paragraph('')
        p.add_run('Understanding the data requires good understanding of the domain and/or access to a '
                'subjectmatter expert (SME) to help make decisions about data quality and data usage:').bold = True

        document.add_paragraph('What are the columns and what do they mean?', style='List Bullet')
        document.add_paragraph('How should we interpret each column and the possible values of a column?', style='List Bullet')
        document.add_paragraph('Should the columns be renamed (and cleaned/trimmed)?', style='List Bullet')
        document.add_paragraph('Are there redundant columns that may share similar information that '
                               'could be dropped in favour of one master column?', style='List Bullet')
        document.add_paragraph('Can columns with no values (or all empty values) be dropped?', style='List Bullet')
        document.add_paragraph('Can columns above a certain threshold of blank values be dropped?', style='List Bullet')
        document.add_paragraph('Can rows that have missing values for certain columns or '
                               'combinations of columns be dropped?', style='List Bullet')
        document.add_paragraph('Can the numeric data type columns be converted / down casted to '
                               'optimise memory usage based on the data values?', style='List Bullet')
        document.add_paragraph('Can some string/object columns be converted to category types?', style='List Bullet')
        document.add_paragraph('Can any columns be discarded that may not be required for analytics?', style='List Bullet')
        document.add_page_break()


    #Page 3a
    if title != True: #If title was user specified, use that user specified title
        document.add_heading(title, 0)
        p = document.add_paragraph(' ')
    document.add_heading('Columns Data Profile Summary', 0)
    p = document.add_paragraph(' ')
    #Heading 1
    document.add_heading('Dataset Shape', level=1)

    #Create a table showing the number of rows and columns
    table = document.add_table(rows=2, cols=3, style = 'Medium Shading 1 Accent 3')
    #Header rows
    cell = table.cell(0, 0)
    cell.text = 'No. of Rows'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True

    cell = table.cell(0, 1)
    cell.text = 'No. of Columns'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    
    cell = table.cell(0, 2)
    cell.text = 'No. of Analyzed Columns'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True

    #Values
    cell = table.cell(1, 0)
    cell.text = F'{df.shape[0] :,}'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = False

    cell = table.cell(1, 1)
    cell.text = str(input_cols)
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = False
    
    cell = table.cell(1, 2)
    cell.text = F'{df.shape[1] :,}'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = False

    #Page 3b
    p = document.add_paragraph(' ')
    #Heading 2
    document.add_heading('Dataframe Columns Summary', level=1)
    p = document.add_paragraph('Note: The min/max values for categories are the min/max lengths of the values. '
                               'Ex: "Alexandria" -> 10, and "Alex" -> 4.')

    #Reshape the column data type dataframe into a form that can be printed in MS Word
    data = round(data_qlt_df[['column_name','data_type', 'unique_values_count', 'min', 'max', 'non_null_values', 'null_values']], 2)
    dataheaderlist = ['Column Name', 'Data Type', '# Unique Values', 'Min', 'Max', 'Non-null Values', 'Null Values']

    #Add a table to the end and create a reference variable
    # Note: the extra added row is so that we can add the header row
    table = document.add_table(data.shape[0]+1, data.shape[1], style='Medium Shading 1 Accent 3')

    #Add the header rows
    for j in range(data.shape[1]):
        #Header row first two columns
        if j <= 1:
            cell = table.cell(0, j)
            cell.text = dataheaderlist[j]
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(9)
            cell_font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            cell = table.cell(0, j)
            cell.text = dataheaderlist[j]
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(9)
            cell_font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cell.width = Inches(1.0)

    #Add the rest of the data frame
    for i in range(data.shape[0]):
        for j in range(data.shape[1]):
            if j <= 1: #First 2 columns are left aligned
                cell = table.cell(i+1, j)
                cell.text = F'{data.values[i,j]}'
                cell_font = cell.paragraphs[0].runs[0].font
                cell_font.size = Pt(9)
                cell_font.bold = False
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif j == 3 or j == 4: #min/max are may return floats 
                cell = table.cell(i+1, j)
                if str(data.values[i,1]) == 'object - datetime': 
                    cell.text = str(data.values[i,j].strftime("%x")) #Only show d/m/y
                elif str(data.values[i,1]) == 'float64': 
                    cell.text = str(str('%.2f' % (data.values[i,j]))) #Show 2 decimal places
                else: #Categories, ints should have integer results and need to be cleaned up
                    cell.text = str(str('%.0f' % (data.values[i,j]))) #No decimal places
                cell_font = cell.paragraphs[0].runs[0].font
                cell_font.size = Pt(9)
                cell_font.bold = False
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else: #Rest are floats
                cell = table.cell(i+1, j)
                cell.text = str(data.values[i,j])
                cell_font = cell.paragraphs[0].runs[0].font
                cell_font.size = Pt(9)
                cell_font.bold = False  
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cell.width = Inches(1.0)

    document.add_page_break()
    

    #Page 4
    #Heading 1
    if not null_vals_df.empty: #Don't print this page if null values df is empty (no largely-empty columns in the df)
        document.add_heading('Columns with Null Values greater than ' + "{:,.2f}".format((1-threshold_perc)*100) + '%', level=1)

        p = document.add_paragraph('The columns should contain at least  ' + "{:,.0f}".format(col_vals_threshold) + 
                                   '  (' + "{:,.2f}".format((col_vals_threshold/df.shape[0])*100) + '%) filled rows out of  '
                                   + "{:,}".format(df.shape[0]) + ' rows to be considered useful.')
        p = document.add_paragraph('Note:  The non-empty values threshold can be set using the threshold_perc '
                                   'variable in the code.')

        #Reshape the column data type dataframe into form that can be printed in MS Word, sort by null counts
        data = round(null_vals_df.sort_values("null_values"), 2)  
        dataheaderlist = ['Column Name', 'Type', '# Unique Values', 'Min Value', 'Max Value', '# Non-Null Values', '# Null Values', '% Null Values'] 

        table = document.add_table(data.shape[0]+1, data.shape[1], style = 'Medium Shading 1 Accent 3')

        #Add the header rows.
        for j in range(data.shape[1]):
            #Header row first columns
            if j <= 1:
                cell = table.cell(0, j)
                cell.text = dataheaderlist[j]
                cell_font = cell.paragraphs[0].runs[0].font
                cell_font.size = Pt(9)
                cell_font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                cell = table.cell(0, j)
                cell.text = dataheaderlist[j]
                cell_font = cell.paragraphs[0].runs[0].font
                cell_font.size = Pt(9)
                cell_font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        #Add the rest of the data frame
        for i in range(data.shape[0]):
            for j in range(data.shape[1]):
                if j <= 1: #Left align the first 2 columns
                    cell = table.cell(i+1, j)
                    cell.text = F'{data.values[i,j]}'
                    cell_font = cell.paragraphs[0].runs[0].font
                    cell_font.size = Pt(9)
                    cell_font.bold = False
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif j == 3 or j == 4: #Min/max can be floats
                    cell = table.cell(i+1, j)
                    if str(data.values[i,1]) == 'object - datetime':
                        cell.text = str(data.values[i,j].strftime("%x")) #Only show d/m/y
                    elif str(data.values[i,1]) == 'float64':
                        cell.text = str(str('%.2f' % (data.values[i,j]))) #Show 2 decimal places
                    else: #Categories, ints should have integer results and need to be cleaned up
                        cell.text = str(str('%.0f' % (data.values[i,j]))) #No decimal places
                    cell_font = cell.paragraphs[0].runs[0].font
                    cell_font.size = Pt(9)
                    cell_font.bold = False
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif j <= 6: #Columns 2, 5, 6 always ints
                    cell = table.cell(i+1, j)
                    cell.text = F'{data.values[i,j]}'
                    cell_font = cell.paragraphs[0].runs[0].font
                    cell_font.size = Pt(9)
                    cell_font.bold = False
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else: #Final column cut to 1 decimal point
                    cell = table.cell(i+1, j)
                    cell.text = F'{data.values[i,j] :,.1f}'
                    cell_font = cell.paragraphs[0].runs[0].font
                    cell_font.size = Pt(9)
                    cell_font.bold = False  
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT  
            
        p = document.add_paragraph(' ')
        if len(nan_fields) > 0:
            p = document.add_paragraph('The following columns were excluded from analysis because they were fully null:')
            p = document.add_paragraph(str(nan_fields))
            p = document.add_paragraph(' ')
        p = document.add_paragraph('Generally columns with extremely large percentages of empty values can be dropped '
                                   'from the dataset as they will not add any value to the analysis.')
        p = document.add_paragraph('')
        p.add_run('However, this depends on the domain of the dataset '
                  'and usage patterns for the columns and data.').bold = True
        document.add_page_break()
                           

    #Page 5: Correlation Plot
    if len(measure_fields) > 1 and 'measure' in subtitle: #Plot will be absent/pointless for 0-1 measure fields
        document.add_heading('Data Correlation Plot (Numeric Columns)', 0)
        p = document.add_paragraph('')

        #Create/plot data correlation.  NOTE:  Only works for numerical data
        du.data_correlation_graph(df[measure_fields]) #Creates the plot for the next line of code
        document.add_picture('./figures/fig_cor_plot.png', height=Inches(6), width=Inches(6))
        
        p = document.add_paragraph(' ')
        p = document.add_paragraph("A data correlation plot indicates how well correlated two different variables are. "
                                  "A value close to 1 indicates that the two are positively correlated.  Ex: The size of a "
                                  "garage (measured in square feet) and the number of cars that the garage can hold. "
                                  "The two variables are positively correlated because as the size of the garage increases, "
                                  "so does the number of cars that can fit in it.  In the plot presented, the diagonal "
                                  "values all read 1 because these entries both read from the same column.  Since the "
                                  "two 'variables' are actually identical, they have a perfect correlation value of 1.")
        p = document.add_paragraph("Correlation values close to -1 are negatively correlated.  These are two variables that "
                                   "behave oppositely to one another.   Ex: Temperature and house heating costs.  As "
                                   "the temperature outside increases, most families will spend less money heating "
                                   "their houses.")
        p = document.add_paragraph("Values close to 0 show neither positive nor negatively correlation.  These variables "
                                   "do not appear to follow similar patterns. ")
        document.add_page_break() 
        
        
    #Page 6: Categorical Column Association Plot
    if len(categorical_fields) > 1 and 'categorical' in subtitle:
        document.add_heading('Categorical Column Associations', 0)
        p = document.add_paragraph('')
        
        #Find categorical columns with more than 1 but also <100 unique values
        catfields_lowuniques = []
        catfields_highuniques = []
        catfields_nouniques = []
        for field in categorical_fields:
            indexno = data_qlt_df.index[data_qlt_df['column_name'] == field]
            if data_qlt_df['unique_values_count'][indexno[0]] < 100 and data_qlt_df['unique_values_count'][indexno[0]] > 1: #If cat field has 2-100 unique values
                catfields_lowuniques.append(field)
            if data_qlt_df['unique_values_count'][indexno[0]] > 100:
                catfields_highuniques.append(field)
            else:
                catfields_nouniques.append(field)
            
        
        #Create/plot categorical column association  NOTE:  Only works for categorical data
        if len(catfields_lowuniques) > 1:
            catcatdf = du.cat_associations(df[catfields_lowuniques], catfields_lowuniques) #Creates the plot for the next line of code
            document.add_picture('./figures/fig_cat_cor_plot.png', height=Inches(6), width=Inches(6))
        
            p = document.add_paragraph("")
            p = document.add_paragraph("Categorical columns plotted above:  " + str(catfields_lowuniques))
            if len(catfields_highuniques) > 0:
                p = document.add_paragraph("")
                p = document.add_paragraph("Categorical columns omitted due to high (>100) unique values:  " + str(catfields_highuniques))
            if len(catfields_nouniques) > 0:
                p = document.add_paragraph("")
                p = document.add_paragraph("Categorical columns omitted due to only having 1 unique value:  " + str(catfields_nouniques))
            p = document.add_paragraph("")
            p = document.add_paragraph("Plotted above are the column associations for categorical columns in the dataframe.  "
                                       "Unlike a numeric correlation plot, values only range from 0 to 1 (rather than -1 to 1).  "
                                       "A value of 0 indicates that knowledge of the category on the bottom provides no information "
                                       "regarding the corresponding category on the left.  A value of 1 indicates that knowlegde of "
                                       "the category on the bottom provides perfect information of the corresponding category on the left.  "
                                       "Columns with values close to 1.0 are likely to be closely related.")
            p = document.add_paragraph("")
            p = document.add_paragraph("Note that this categorical column association plot is not symmetric like a numeric "
                                       "correlation plot.  Instead, the plot shows to what extent knowledge of the category on the left "
                                       "provides information regarding the corresponding category on the bottom of the plot. "
                                       "An example of the difference relating two categories can be demonstrated by considering bees "
                                       "and wasps.  Both fall into different families of insects (bees vs. wasps), but have the same color scheme "
                                       "(black/yellow).  Thus, while knowing the insect family is bee or wasp gives perfect information "
                                       "about the insect's color, knowledge of the color does not provide perfect information about "
                                       "the insect's family (knowing the color is black/yellow only narrows the family down to either "
                                       "bee or wasp).")     
            document.add_page_break()
            
        else:
            p = document.add_paragraph("Categorical correlations not determined due to excessive (>100) unique values in at least one category "
                                       "in each category-category pair.")
            document.add_page_break()
        
        
    #Page 7: Categorical Column vs Numeric Column Association Plot
    if len(catfields_lowuniques) > 1 and 'categorical' in subtitle and len(measure_fields) > 1 and 'measure' in subtitle:
        document.add_heading('Categorical Column and Measure Column Associations', 0)
        p = document.add_paragraph('')
        
        #Create/plot the measure/categorical column associations  NOTE: Needs both categorical and measure data
        catnumdf = du.catnum_associations(df, catfields_lowuniques, measure_fields)  #Creates the plot for the next line of code
        document.add_picture('./figures/fig_catnum_cor_plot.png', height=Inches(6), width=Inches(6))
        
        p = document.add_paragraph("")
        p = document.add_paragraph("Categorical columns plotted above:  " + str(catfields_lowuniques))
        p = document.add_paragraph("")
        p = document.add_paragraph("Measure columns plotted above:  " + str(measure_fields))
        p = document.add_paragraph("")
        p = document.add_paragraph("Plotted above are the column associations for categorical columns with numeric columns in the dataframe.  "
                                   "Values range from 0 to 1.  Numeric columns are shown on the x axis and categorical columns are shown "
                                   "on the y axis.  A value of 0 indicates that knowledge of the measure value provides no "
                                   "information regarding the categorical variable.  A value of 1.0 indicates that knowledge of the measure "
                                   "value allows for perfect identification of the categorical variable.  Columns with values close to 1.0 "
                                   "are likely to be closely related.")        
        document.add_page_break()
                                      
                           
    #Page 8+:  Summarize data for each column
    document.add_heading('Column Data Profile Details', 0)
    
    #Summarize plots that will be shown, using inputs from subtitle to determine which summaries are shown
    if 'measure' in subtitle:
        p = document.add_paragraph("")
        p.add_run('Measure Data Plots').bold = True
        p = document.add_paragraph("Each measure (numerical) column is summarized with six plots:")
        p = document.add_paragraph("1. Box plot for all values.  This plot shows the median value with an orange line."
                                   "  The middle half of the data (25th to 75 percentile) falls within the box that is "
                                   "bisected by the orange median line. The 'whiskers' above and below the box extend "
                                   "to either the minimum or maximum values or to 1.5 times the difference between the top "
                                   "and bottom values of the box, if values still exist beyond this range.  Any outlier "
                                   "values that exist even beyond the reach of the whiskers are shown as small dots.  "
                                   "The box plot helps by showing the distribution of the data around the mean and by "
                                   " highlighting any outliers.")
        p = document.add_paragraph("2. Distribution plot.  This plot is a simple visualization of the distribution of the "
                                   "data.  Each column shows the number of values that fall within the width of the column.")
        p = document.add_paragraph("3. Box plot without outliers.  This plot is similar to the first, but with all outliers "
                                   "(values existing beyond the reach of the whiskers) removed to improve the visibility of "
                                   "the remainder of the distribution.")
        p = document.add_paragraph("4. Violin plot (for values <95th percentile).  A violin plot is similar to a box plot in "
                                   "that it is also used for visualizing the distribution of the data.  The wider the "
                                   "distribution is at a given value, the more often the column has results around that value.")
        p = document.add_paragraph("5. Quartile box plot.  This plot is similar to the original box plot, except that each "
                                   "quartile of data (0-25th,25th-50th,50th-75th,75th-100) is individually plotted as its own "
                                   "box plot.  This plot gives better insight into the distribution of each quartile of data.  "
                                   "If less than 4 box plots appear here, it means that the data is not evenly distributed, "
                                   "and that multiple boxes would share identical shapes.  Ex:  A column composed of 100 "
                                   "1's and a single 0.  In this case, three of the quartiles would be identical, as "
                                   "they would be composed entirely of 25 1's.  In the case that less than 4 boxes appear, "
                                   "it can be implied the missing boxes would exactly match the shape and size of one of "
                                   "the existing boxes.")
        p = document.add_paragraph("6. Distribution plot (for values <95th percentile).  This plot is similar to the previous "
                                   "distribution plot, but it ignores data above the 95th percentile.  Additionally, it "
                                   "highlights which columns fall within which percentile, based on the value of the left "
                                   "side of the column. (Ex: If a column encompasses the entire 0-25th percentile as well "
                                   "as part of the 25th-75th percentile, it will be colored yellow because the left side "
                                   "of the column would be the 0th percentile, and thus the whole column is grouped under "
                                   "this coloring.)")
        p = document.add_paragraph(" ")
    if 'date' in subtitle:
        p = document.add_paragraph("")
        p.add_run('Date Data Plots').bold = True
        p = document.add_paragraph("Each date column is summarized with a plot condensing each date down to the year and month "
                                   "that the date is recorded under.  Each plot indicates the number of dates that were "
                                   "collected for that year and month combination.")
        p = document.add_paragraph(" ")
    if 'categorical' in subtitle:
        p = document.add_paragraph("")  
        p.add_run('Categorical Data Plots').bold = True
        p = document.add_paragraph("Each categorical column is summarized with a plot that lists the most common "
                                   "values that appear within the column, as well as the number of apppearances "
                                   " of each of these values.")
        p = document.add_paragraph("Note: The min/max/mean values in the category tables are based on the length of "
                                   "the categories as strings.  Ex: 'Alexandria' -> 10, so if the longest entry into "
                                   "a category column was 'Alexandria', the max value for this category would be 10.")
    document.add_page_break()
    
    #Sort dataframe by type ('categorical','date','measure') so same typed columns all appear in sequence together
    data_qlt_df = data_qlt_df.sort_values(['input_type'], ascending = [False]) #'measure' 1st, 'date' 2nd, 'categorical' 3rd
    data_qlt_df = data_qlt_df.reset_index(drop=True)
    
    for ind in range(data_qlt_df.shape[0]): #Make a page for each included column in df
        #Use subtitle argument from function definition to determine if plots should be shown
        if ('date' in subtitle and data_qlt_df['column_name'][ind] in date_fields) or ('categorical' in subtitle and data_qlt_df['column_name'][ind] in categorical_fields) or ('measure' in subtitle and data_qlt_df['column_name'][ind] in measure_fields):
            
            #Creating a table for a date or category column
            if data_qlt_df['column_name'][ind] in date_fields or data_qlt_df['column_name'][ind] in categorical_fields:
                tablerows = 4 #Include min/mean/max row 
            #Creating a table for a measure column
            if data_qlt_df['column_name'][ind] in measure_fields: 
                tablerows = 5 #Also include 25/50/75 percentile row (in addition to min/mean/max row)     

            #Create table for column profile details
            table = document.add_table(rows=tablerows, cols=6, style = 'Medium Shading 1 Accent 3' )

            #Merge cells in header row for Column Name
            for y in range(len(table.rows[0].cells)-1):
                a = table.cell(0,y)
                b = table.cell(0,y+1)
                a.merge(b)

            #Merge cells in detail rows spanning 2 cells x 3 
            for row in range(1,tablerows):
                a = table.cell(row,0)
                b = table.cell(row,1)
                a.merge(b)
                a = table.cell(row,2)
                b = table.cell(row,3)
                a.merge(b)
                a = table.cell(row,4)
                b = table.cell(row,5)
                a.merge(b)


            #ADD VALUES TO TABLE#
            #Cell 0,0 (merged 6 cells): Header - Column Name
            cell = table.cell(0, 0)
            cell.text = data_qlt_df["column_name"][ind]
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(15)
            cell_font.bold = True

            #Cell 1,0: Blank
            cell = table.cell(1, 1)
            cell.text = "Column Name :\n"
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            cell_font.bold = True
            p = cell.paragraphs[0].add_run(str(data_qlt_df["column_name"][ind]))
            cell_font2 = cell.paragraphs[0].runs[1].font
            cell_font2.size = Pt(12)
            cell_font2.bold = False

            #Cell 1,1: Column data type
            cell = table.cell(1, 3)
            cell.text = 'Data Type : \n'
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            cell_font.bold = True
            p = cell.paragraphs[0].add_run(str(data_qlt_df["data_type"][ind]))
            cell_font2 = cell.paragraphs[0].runs[1].font
            cell_font2.size = Pt(12)
            cell_font2.bold = False

            #Cell 1,2: Count of total values in the column
            cell = table.cell(1, 5)
            cell.text = 'Values Count : \n'
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            cell_font.bold = True
            p = cell.paragraphs[0].add_run(F'{data_qlt_df["count"][ind] :,.0f}')
            cell_font2 = cell.paragraphs[0].runs[1].font
            cell_font2.size = Pt(11)
            cell_font2.bold = False

            #Cell 2,0: Count of unique values in the column
            cell = table.cell(2, 1)
            cell.text = 'Unique Values Count : \n'
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            cell_font.bold = True
            unique_per = (data_qlt_df["unique_values_count"][ind] / data_qlt_df["count"][ind]) * 100
            p = cell.paragraphs[0].add_run(F'{data_qlt_df["unique_values_count"][ind] :,.0f}' + "   " + 
                                           F'({unique_per :,.2f}%)' )
            cell_font2 = cell.paragraphs[0].runs[1].font
            cell_font2.size = Pt(11)
            cell_font2.bold = False

            #Cell 2,1: Count of non-null values in the column
            cell = table.cell(2, 3)
            cell.text = 'Non-Null Values Count : \n'
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            cell_font.bold = True
            p = cell.paragraphs[0].add_run(F'{data_qlt_df["non_null_values"][ind] :,.0f}' + "   " + 
                                           F' ({data_qlt_df["%_of_non_nulls"][ind]  :,.2f}%)' )
            cell_font2 = cell.paragraphs[0].runs[1].font
            cell_font2.size = Pt(11)
            cell_font2.bold = False       

            #Cell 2,2: Count of null values in the column
            cell = table.cell(2, 5)
            cell.text = 'Null Values Count : \n'
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            cell_font.bold = True
            p = cell.paragraphs[0].add_run(F'{data_qlt_df["null_values"][ind]  :,.0f}' + "   " + 
                                           F' ({data_qlt_df["%_of_nulls"][ind]  :,.2f}%)' )
            cell_font2 = cell.paragraphs[0].runs[1].font
            cell_font2.size = Pt(11)
            cell_font2.bold = False

            #Cell 3,0: Min of values in the column
            cell = table.cell(3, 1)
            cell.text = 'Min : \n'
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            cell_font.bold = True
          
            if tablerows == 4: #If grabbing a date/category min
                if str(data_qlt_df["data_type"][ind]) == 'object' or str(data_qlt_df["data_type"][ind]) == 'category':
                    p = cell.paragraphs[0].add_run(str('%.0f' % (data_qlt_df["min"][ind])))   
                else: #If a datetime
                    p = cell.paragraphs[0].add_run(str(data_qlt_df["min"][ind]))   
            if tablerows == 5: #If grabbing a measure min
                if str(data_qlt_df["data_type"][ind]) == 'int64':
                    p = cell.paragraphs[0].add_run(str('%.0f' % (data_qlt_df["min"][ind]))) 
                else: #Leave 2 decimal places on floats
                    p = cell.paragraphs[0].add_run(F'{data_qlt_df["min"][ind]  :,.2f}' )
            cell_font2 = cell.paragraphs[0].runs[1].font
            cell_font2.size = Pt(11)
            cell_font2.bold = False

            #Cell 3,1: Mean of values in the column
            cell = table.cell(3, 3)
            cell.text = 'Mean :  \n'
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            cell_font.bold = True
            if tablerows == 4: #Date/category mean
                p = cell.paragraphs[0].add_run(str(data_qlt_df["mean"][ind]))
            if tablerows == 5: #Measure mean
                p = cell.paragraphs[0].add_run(F'{data_qlt_df["mean"][ind] :,.2f}' )
            cell_font2 = cell.paragraphs[0].runs[1].font
            cell_font2.size = Pt(11)
            cell_font2.bold = False

            #Cell 3,2: Max of values in the column
            cell = table.cell(3, 5)
            cell.text = 'Max : \n'
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            cell_font.bold = True
            if tablerows == 4: #If grabbing a date/category max
                if str(data_qlt_df["data_type"][ind]) == 'object' or str(data_qlt_df["data_type"][ind]) == 'category':
                    p = cell.paragraphs[0].add_run(str('%.0f' % (data_qlt_df["max"][ind])))   
                else:
                    p = cell.paragraphs[0].add_run(str(data_qlt_df["max"][ind])) 
                #p = cell.paragraphs[0].add_run(str(data_qlt_df["max"][ind]))                           
            if tablerows == 5: #If grabbing a measure max
                if str(data_qlt_df["data_type"][ind]) == 'int64':
                    p = cell.paragraphs[0].add_run(str('%.0f' % (data_qlt_df["max"][ind]))) 
                else: #Leave 2 decimal places on floats
                    p = cell.paragraphs[0].add_run(F'{data_qlt_df["max"][ind]  :,.2f}' )
                #p = cell.paragraphs[0].add_run(F'{data_qlt_df["max"][ind]  :,.2f}' )
            cell_font2 = cell.paragraphs[0].runs[1].font
            cell_font2.size = Pt(11)
            cell_font2.bold = False

            if tablerows == 5: #If table is for measures
                #Cell 4,0: 25th Percentile of values in the column
                cell = table.cell(4, 1)
                cell.text = '25th Percentile : \n'
                cell_font = cell.paragraphs[0].runs[0].font
                cell_font.size = Pt(11)
                cell_font.bold = True
                p = cell.paragraphs[0].add_run(F'{data_qlt_df["25%"][ind]  :,.2f}' )
                cell_font2 = cell.paragraphs[0].runs[1].font
                cell_font2.size = Pt(11)
                cell_font2.bold = False

                #Cell 4,1: 50th Percentile of values in the column
                cell = table.cell(4, 3)
                cell.text = '50th Percentile : \n'
                cell_font = cell.paragraphs[0].runs[0].font
                cell_font.size = Pt(11)
                cell_font.bold = True
                p = cell.paragraphs[0].add_run(F'{data_qlt_df["50%"][ind]  :,.2f}' )
                cell_font2 = cell.paragraphs[0].runs[1].font
                cell_font2.size = Pt(11)
                cell_font2.bold = False

                #Cell 4,2: 75th Percentile of values in the column
                cell = table.cell(4, 5)
                cell.text = '75th Percentile : \n'
                cell_font = cell.paragraphs[0].runs[0].font
                cell_font.size = Pt(11)
                cell_font.bold = True
                p = cell.paragraphs[0].add_run(F'{data_qlt_df["75%"][ind]  :,.2f}' )
                cell_font2 = cell.paragraphs[0].runs[1].font
                cell_font2.size = Pt(11)
                cell_font2.bold = False

            p = document.add_paragraph(' ')

            # Dates/Categories have wide plots
            if data_qlt_df['column_name'][ind] in date_fields or data_qlt_df['column_name'][ind] in categorical_fields: 
                try:
                    document.add_picture('./figures/' + str(data_qlt_df['column_name'][ind]) + '.png', height=Inches(6.0), width=Inches(7))
                except:
                    document.add_paragraph('Image not found')     
            # Measures have different plot names + tall plots
            elif data_qlt_df['column_name'][ind] in measure_fields: 
                try:
                    document.add_picture('./figures/' + str(data_qlt_df['column_name'][ind] + '.png'), height=Inches(6.5), width=Inches(7.0))
                except:
                    document.add_paragraph('Image not found')  

            if ind != data_qlt_df.shape[0]-1: #Add page break on all except last page
                document.add_page_break()
        
        else: #Plot is not encompassed by subtitle arguments and will not be generated
            pass
        
    #Page 9+:  Correlation plots
    if corrplots == True and len(measure_pair_list) > 0:
        document.add_heading('Correlation Plots (Values within 3 Std. Dev.)', 0)
        du.data_correlation_graph(df[measure_fields]) #Creates the plot for the next line of code
        document.add_picture('./figures/fig_cor_plot.png', height=Inches(6), width=Inches(6))
        p = document.add_paragraph("")
        p = document.add_paragraph("Each measure (numerical) column's correlations can be measured as shown in the "
                                   "data correlation plot. A value close to 1 indicates that the two are positively correlated. "
                                   "A column with the values 1,2,3 would be perfectly correlated (have a value of 1) with a "
                                   "column that has values 2,4,6.  The second column's values are directly relatable to the "
                                   "first column's values (column 2 = column 1 * 2).  A second column with values -2,-4,-6 would "
                                   "be perfectly inversely correlated with column one and have a correlation value of -1.  "
                                   "Plots showing these relationships are shown for all columns with correlation values "
                                   "greater than 0.5 or lower than -0.5.  Only values within 3 standard deviations of the mean "
                                   "are shown in order to prevent distortion from extreme outliers.")  
        document.add_page_break()

        for pair in measure_pair_list: #Look at all combinations of two measure values
            corrtable = df[pair].corr(method = 'pearson')
            corrvalue = corrtable[pair[0]][1]
            document.add_heading('Columns ' + str(pair[0]) + ' and ' + str(pair[1]) + ':  Correlation = %.2f' % corrvalue, level=1)
            try:
                document.add_picture('./figures/' + str(pair[0]) + 'vs' + str(pair[1]) + '.png', height=Inches(6.5), width=Inches(7))
            except:
                document.add_paragraph('Image not found') 
                
            document.add_page_break()  
            
            
    #Page 10+:  Categorical-Categorical plots
    #First ensure there's categorical-categorical relationships to plot
    if len(categorical_fields) >= 2:
        
        #Count how many categorical-categorical relationships are greater than our threshold of 0.75; if none, don't make this page
        plotcounter = 0
        catcatlist = [] #List of pairs to plot
        for i, j in catcatdf.iterrows(): #i tracks the plot index
            for columns in catcatdf:
                if catcatdf[columns][i] > 0.75 and columns != i: #Check for threshold and that both cols aren't the same
                    plotcounter = plotcounter + 1
                    catcatlist.append([i,columns])
        
        if plotcounter > 0: #If at least one categorical-categorical pair has high correlation, we'll make these plots
            document.add_heading('Categorical-Categorical Correlation Plots', 0)
            p = document.add_paragraph("")
            p = document.add_paragraph("Plotted here are the heatmaps showing the frequency with which any row in the dataframe has "
                                       "the values shown on the x and y axis for the two respective columns.  Intersections that appear "
                                       "dark blue show frequent overlap between the two column values.")
            p = document.add_paragraph("Note:  The plots shown only show the top 15 values in each categorical column in order "
                                       "to prevent columns with many different values from producing extremely large plots.  "
                                       "Plots showing these relationships are shown for all columns with correlation values "
                                       "greater than 0.75.")
            document.add_page_break()
            
            #Plot each correlation with value greater than 0.75 as found above:
            for pair in catcatlist:
                corrvalue = catcatdf[pair[1]][pair[0]]
                document.add_heading('Columns ' + str(pair[0]) + ' and ' + str(pair[1]) + ':  Correlation = %.2f' % corrvalue, level=1)
                du.cat_cat_heatmap(df, pair) #Create the heatmap
                try:
                    document.add_picture('./figures/' + str(pair[0]) + '_vs_' + str(pair[1]) + '_Heatmap.png', height=Inches(6.5), width=Inches(7))
                except:
                    try:
                        document.add_picture('./figures/' + str(pair[1]) + '_vs_' + str(pair[0]) + '_Heatmap.png', height=Inches(6.5), width=Inches(7))
                    except:
                        document.add_paragraph('Image not found')
                    
                document.add_page_break()
                
                        
    #Page 11+:  Categorical-Numeric plots
    #First ensure there's categorical-numeric relationships to plot
    if len(categorical_fields) >=1 and len(measure_fields) >=1:
        
        #Count how many categorical-numeric relationships are greater than our threshold of 0.75; if none, don't make this page
        plotcounter = 0
        catnumlist = [] #List of pairs to plot
        for i, j in catnumdf.iterrows(): #i tracks the plot index
            for columns in catnumdf:
                if catnumdf[columns][i] > 0.75:  #Check for threshold
                    plotcounter = plotcounter + 1
                    catnumlist.append([i,columns])
                
        if plotcounter > 0: #If at least one categorical-numeric pair has a high correlation, we'll make these plots
            document.add_heading('Categorical-Numeric Correlation Plots', 0)
            p = document.add_paragraph("")
            p = document.add_paragraph("Plotted here are the values of a numeric column corresponding to the values in a categorical column.  "
                                       "Plots showing these relationships are shown for all categorical-numeric combinations with correlation "
                                       "values greater than 0.75.  Combinations with high correlation values demonstrate that knowledge of "
                                       "the continuous variable (the numeric variable) allows you to predict the value of the categorical "
                                       "value with high probability, indicating that the two columns are closely related.")
            p = document.add_paragraph("The values in the numeric column are shown for each of the 15 most common values in the categorical column as a box plot "
                                       "without outliers.  The median value is shown as a green line, and the endcaps of the box represent "
                                       "the locations of the 25th and 75th percentiles, thus making the box encapsulate the middle "
                                       "50% of data from each column value.  The whiskers extend to the largest or lowest values, "
                                       "unless there are outliers outside of this range.  In that case, they extend to 1.5 times the "
                                       "difference between the 75th and 25th percentiles.")
            p = document.add_page_break()
            
            #Plot each correlation with value greater than 0.75 as found above:
            for pair in catnumlist:
                corrvalue = catnumdf[pair[1]][pair[0]]
                document.add_heading('Columns ' + str(pair[0]) + ' and ' + str(pair[1]) + ':  Correlation = %.2f' % corrvalue, level=1)
                du.cat_num_heatmap(df, pair) #Create the heatmap
                try:
                    document.add_picture('./figures/' + str(pair[1]) + '_vs_' + str(pair[0]) + '_Plot.png', height=Inches(6.5), width=Inches(7))
                except:
                    try:
                        document.add_picture('./figures/' + str(pair[0]) + '_vs_' + str(pair[1]) + '_Plot.png', height=Inches(6.5), width=Inches(7))
                    except:
                        document.add_paragraph('Image not found')
                
                document.add_page_break()
                
    
    #Page 12:  Summary Page
    #Concisely list all columns that have many nulls (25%) as well as % null
    # Also list all heavily correlated columns, as well as correlation values
    
    document.add_heading('DataFrame Summary', 0)
    document.add_heading('Summary: Columns with Null Values >' + "{:,.2f}".format((1-threshold_perc)*100) + '%', level=1)
    
    if null_vals_df.empty: 
        document.add_paragraph("No columns in the submitted dataframe were heavily null.")
    else:
        data = round(null_vals_df.sort_values("null_values"), 2)  
        data = data.iloc[:, [0] + [-1]] #Only grab fist and last column of null_vals_df (column name and % null)
        dataheaderlist = ['Column Name', '% Null Values'] 

        table = document.add_table(data.shape[0]+1, data.shape[1], style = 'Medium Shading 1 Accent 3')

        #Add the header rows.
        for j in range(data.shape[1]):
            #Header row first columns
            if j < 1: #Left align first column (column names)
                cell = table.cell(0, j)
                cell.text = dataheaderlist[j]
                cell_font = cell.paragraphs[0].runs[0].font
                cell_font.size = Pt(14)
                cell_font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            else: #Right align second column (% null)
                cell = table.cell(0, j)
                cell.text = dataheaderlist[j]
                cell_font = cell.paragraphs[0].runs[0].font
                cell_font.size = Pt(14)
                cell_font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        #Add the rest of the data frame
        for i in range(data.shape[0]):
            for j in range(data.shape[1]):
                if j < 1: #Left align first column (column names)
                    cell = table.cell(i+1, j)
                    cell.text = F'{data.values[i,j]}'
                    cell_font = cell.paragraphs[0].runs[0].font
                    cell_font.size = Pt(12)
                    cell_font.bold = False
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                else: #Right align second column, cut to 1 decimal point
                    cell = table.cell(i+1, j)
                    cell.text = F'{data.values[i,j] :,.1f}'
                    cell_font = cell.paragraphs[0].runs[0].font
                    cell_font.size = Pt(12)
                    cell_font.bold = False  
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT  
            
        p = document.add_paragraph(' ')
        if len(nan_fields) > 0:
            p = document.add_paragraph('The following columns were excluded from analysis because they were fully null:')
            p = document.add_paragraph(str(nan_fields))
            p = document.add_paragraph(' ')
        p = document.add_paragraph('Generally columns with extremely large percentages of empty values can be dropped '
                                   'from the dataset as they will not add any value to the analysis.')
        document.add_page_break()
    
    
    #Numerical-Numerical associations
    if len(measure_fields) > 1 and corrplots == True: #Only attempt if associations are possible
        document.add_heading('Strongly Correlated Measure-Measure Column Pairs (>0.5, <-0.5)', level=1)
        
        if len(measure_pair_list) == 0: #If no highly-correlated or anti-correlated measure columns
            document.add_paragraph("No measure column pairs were highly correlated or highly anti-correlated.")
        
        else:
            #Create lists of correlated columns and the correlation values
            corrlist = []
            corrvals = []
            for pair in measure_pair_list: #Look at all combinations of two measure values
                corrtable = df[pair].corr(method = 'pearson')
                corrvalue = corrtable[pair[0]][1]
                
                corrlist.append(pair)
                corrvals.append(corrvalue)
            
            #Split list of pairs into two lists
            corrlist1, corrlist2 = map(list, zip(*corrlist))
            #Create new df only containing the pairs with high correlation values, as well as the correlation values
            corrdf = pd.DataFrame(list(zip(corrlist1,corrlist2,corrvals)), columns = ['Correlated Column 1', 'Correlated Column 2', 'Correlation Value'])
            corrdf = corrdf.sort_values(by=['Correlation Value'], ascending=False) #Show values from 1 to -1
            data = corrdf
            
            #Use the new df to create a summary table
            dataheaderlist = ['Correlated Column 1', 'Correlated Column 2', 'Correlation Value'] 
            table = document.add_table(data.shape[0]+1, data.shape[1], style = 'Medium Shading 1 Accent 3')
            
            #Add the header rows.
            for j in range(data.shape[1]):
                #Header row first columns
                if j <= 1: #Left align first two columns (column names)
                    cell = table.cell(0, j)
                    cell.text = dataheaderlist[j]
                    cell_font = cell.paragraphs[0].runs[0].font
                    cell_font.size = Pt(14)
                    cell_font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                else: #Right align second column (Correlation Value)
                    cell = table.cell(0, j)
                    cell.text = dataheaderlist[j]
                    cell_font = cell.paragraphs[0].runs[0].font
                    cell_font.size = Pt(14)
                    cell_font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
            #Add the rest of the data frame
            for i in range(data.shape[0]):
                for j in range(data.shape[1]):
                    if j <= 1: #Left align first two column (column names)
                        cell = table.cell(i+1, j)
                        cell.text = F'{data.values[i,j]}'
                        cell_font = cell.paragraphs[0].runs[0].font
                        cell_font.size = Pt(12)
                        cell_font.bold = False
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else: #Right align second column, cut to 2 decimal points
                        cell = table.cell(i+1, j)
                        cell.text = F'{data.values[i,j] :,.2f}'
                        cell_font = cell.paragraphs[0].runs[0].font
                        cell_font.size = Pt(12)
                        cell_font.bold = False  
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT 
                        
            p = document.add_paragraph(' ')
            p = document.add_paragraph("Data correlation values in numeric variables indicate how well correlated the two variables are. "
                                  "A value close to 1 indicates that the two are positively correlated.  Ex: The size of a "
                                  "garage (measured in square feet) and the number of cars that the garage can hold. "
                                  "The two variables are positively correlated because as the size of the garage increases, "
                                  "so does the number of cars that can fit in it. ")
            p = document.add_paragraph("Correlation values close to -1 are negatively correlated.  These are two variables that "
                                   "behave oppositely to one another.   Ex: Temperature and house heating costs.  As "
                                   "the temperature outside increases, most families will spend less money heating "
                                   "their houses.")
            p = document.add_paragraph("Values close to 0 show neither positive nor negatively correlation.  These variables "
                                   "do not appear to follow similar patterns. ")
            
        document.add_page_break()
    
    
    #Categorical-Categorical associations
    if 'catcatdf' in locals() and corrplots == True: #Check if categorical-categorical associations were determined
        document.add_heading('Strongly Correlated Categorical-Categorical Column Pairs (>0.75)', level=1)
        
        #Create list of correlated columns
        catcatlist = []
        catcatvalue = []
        for i, j in catcatdf.iterrows(): #i tracks the plot index
            for columns in catcatdf:
                if catcatdf[columns][i] > 0.75 and columns != i: #Check for threshold and that both cols aren't the same
                    catcatlist.append([i,columns]) #List of pairs
                    catcatvalue.append(catcatdf[columns][i]) #List of corr values for the pairs
        
        if len(catcatlist)==0: #If no highly correlated categorical-categorical pairs
            document.add_paragraph("No categorical-categorical pairs were heavily correlated.")
            
        else:
            #Split list of pairs into two lists
            catcatlist1, catcatlist2 = map(list, zip(*catcatlist))
            #Create new df only containing the pairs with high correlation values, as well as the correlation values
            catcatcorrdf = pd.DataFrame(list(zip(catcatlist1,catcatlist2,catcatvalue)), columns = ['Correlated Column 1', 'Correlated Column 2', 'Correlation Value'])
            catcatcorrdf = catcatcorrdf.sort_values(by=['Correlation Value'], ascending=False) #Show most correlated columns first
            data = catcatcorrdf
        
            #Use the new df to create a summary table
            dataheaderlist = ['Correlated Column 1', 'Correlated Column 2', 'Correlation Value'] 
            table = document.add_table(data.shape[0]+1, data.shape[1], style = 'Medium Shading 1 Accent 3')
        
            #Add the header rows.
            for j in range(data.shape[1]):
                #Header row first columns
                if j <= 1: #Left align first two columns (column names)
                    cell = table.cell(0, j)
                    cell.text = dataheaderlist[j]
                    cell_font = cell.paragraphs[0].runs[0].font
                    cell_font.size = Pt(14)
                    cell_font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                else: #Right align second column (Correlation Value)
                    cell = table.cell(0, j)
                    cell.text = dataheaderlist[j]
                    cell_font = cell.paragraphs[0].runs[0].font
                    cell_font.size = Pt(14)
                    cell_font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
            #Add the rest of the data frame
            for i in range(data.shape[0]):
                for j in range(data.shape[1]):
                    if j <= 1: #Left align first two column (column names)
                        cell = table.cell(i+1, j)
                        cell.text = F'{data.values[i,j]}'
                        cell_font = cell.paragraphs[0].runs[0].font
                        cell_font.size = Pt(12)
                        cell_font.bold = False
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else: #Right align second column, cut to 2 decimal points
                        cell = table.cell(i+1, j)
                        cell.text = F'{data.values[i,j] :,.2f}'
                        cell_font = cell.paragraphs[0].runs[0].font
                        cell_font.size = Pt(12)
                        cell_font.bold = False  
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT 
                        
            p = document.add_paragraph(' ')
            p = document.add_paragraph("The table above summarizes the column associations for categorical columns in the dataframe.  "
                                       "Values range from 0 to 1.  "
                                       "A value of 0 indicates that knowledge of the second category provides no information "
                                       "regarding the first category.  A value of 1 indicates that knowlegde of "
                                       "the second category provides perfect information of the first category.  "
                                       "Columns with values close to 1.0 are closely related.")
            p = document.add_paragraph("")
            p = document.add_paragraph("Note that this categorical column association plot is not symmetric like a numeric "
                                       "correlation plot.  Instead, the plot shows to what extent knowledge of the category on the left "
                                       "provides information regarding the corresponding category on the bottom of the plot. "
                                       "An example of the difference relating two categories can be demonstrated by considering bees "
                                       "and wasps.  Both fall into different families of insects (bees vs. wasps), but have the same color scheme "
                                       "(black/yellow).  Thus, while knowing the insect family is bee or wasp gives perfect information "
                                       "about the insect's color, knowledge of the color does not provide perfect information about "
                                       "the insect's family (knowing the color is black/yellow only narrows the family down to either "
                                       "bee or wasp).")  
        document.add_page_break()
    
    
    #Categorical-Numeric associations
    if 'catnumdf' in locals() and corrplots == True: #Check if categorical-categorical associations were determined
        document.add_heading('Strongly Correlated Categorical-Measure Column Pairs (>0.75)', level=1)
        
        #Create list of correlated columns
        catnumlist = []
        catnumvalue = []
        for i, j in catnumdf.iterrows(): #i tracks the plot index
            for columns in catnumdf:
                if catnumdf[columns][i] > 0.75: #Check for threshold
                    catnumlist.append([i,columns]) #List of pairs
                    catnumvalue.append(catnumdf[columns][i]) #List of corr values for the pairs
        
        if len(catnumlist)==0: #If no highly correlated categorical-categorical pairs
            document.add_paragraph("No categorical-measure pairs were heavily correlated.")
            
        else:
            #Split list of pairs into two lists
            catnumlist1, catnumlist2 = map(list, zip(*catnumlist))
            #Create new df only containing the pairs with high correlation values, as well as the correlation values
            catnumcorrdf = pd.DataFrame(list(zip(catnumlist1,catnumlist2,catnumvalue)), columns = ['Correlated Column 1', 'Correlated Column 2', 'Correlation Value'])
            catnumcorrdf = catnumcorrdf.sort_values(by=['Correlation Value'], ascending=False) #Show most correlated columns first
            data = catnumcorrdf
        
            #Use the new df to create a summary table
            dataheaderlist = ['Correlated Category', 'Correlated Measure', 'Correlation Value'] 
            table = document.add_table(data.shape[0]+1, data.shape[1], style = 'Medium Shading 1 Accent 3')
            
            #Add the header rows.
            for j in range(data.shape[1]):
                #Header row first columns
                if j <= 1: #Left align first two columns (column names)
                    cell = table.cell(0, j)
                    cell.text = dataheaderlist[j]
                    cell_font = cell.paragraphs[0].runs[0].font
                    cell_font.size = Pt(14)
                    cell_font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                else: #Right align second column (Correlation Value)
                    cell = table.cell(0, j)
                    cell.text = dataheaderlist[j]
                    cell_font = cell.paragraphs[0].runs[0].font
                    cell_font.size = Pt(14)
                    cell_font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
            #Add the rest of the data frame
            for i in range(data.shape[0]):
                for j in range(data.shape[1]):
                    if j <= 1: #Left align first two column (column names)
                        cell = table.cell(i+1, j)
                        cell.text = F'{data.values[i,j]}'
                        cell_font = cell.paragraphs[0].runs[0].font
                        cell_font.size = Pt(12)
                        cell_font.bold = False
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else: #Right align second column, cut to 2 decimal points
                        cell = table.cell(i+1, j)
                        cell.text = F'{data.values[i,j] :,.2f}'
                        cell_font = cell.paragraphs[0].runs[0].font
                        cell_font.size = Pt(12)
                        cell_font.bold = False  
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT 
                        
            p = document.add_paragraph(' ')
            p = document.add_paragraph("The table above summarizes the column associations for categorical columns and "
                                       "measure columns in the dataframe.  Values range from 0 to 1.  A value of 0 indicates that knowledge of the measure value provides no "
                                       "information regarding the categorical variable.  A value of 1.0 indicates that knowledge of the measure "
                                       "value allows for perfect identification of the categorical variable.  Columns with values close to 1.0 "
                                       "are likely to be closely related.")  
        document.add_page_break()
            
     
    
    #Page 13+:  Categorical binary plots
#    if (binaryplots == True) and (len(binarycatlist) >= 2):
#        document.add_heading('Categorical Binary Plots', 0)
#        p = document.add_paragraph("")
#        p = document.add_paragraph("Plotted here are heatmaps showing the frequency with which any row in the dataframe has "
#                                   "the values shown on the x and y axis for the two respective columns.  Intersections that "
#                                   "appear dark blue show frequent overlap between the two column values.")
#        p = document.add_paragraph("Note:  The plots shown are only for categorical columns with 15 or less unique values.  "
#                                   "Additionally, the plots are only a sample of the full original dataframe (in order to "
#                                   "greatly speed up the runtime of the program).")
#        
#        binary_pair_list = []
#        for combo in combinations(binarycatlist, 2): #Look at all combinations of two measure values
#            
#            #Create a list with the unique combination of two fields
#            binary_pair = []
#            for value in combo:
#                binary_pair.append(value)
#            
#            #Append combo to list of lists containing each correlated pair
#            binary_pair_list.append(binary_pair)
#        
#        #Go through our list of lists and grab the corresponding figure for each combination
#        for pair in binary_pair_list:
#            if os.path.exists('./figures/' + str(pair[0]) + '_vs_' + str(pair[1]) + '_Heatmap.png') == True:
#                document.add_heading(str(pair[0]) + " vs. " + str(pair[1]))
#                document.add_picture('./figures/' + str(pair[0]) + '_vs_' + str(pair[1]) + '_Heatmap.png', height=Inches(6.5), width=Inches(7))  
#            elif os.path.exists('./figures/' + str(pair[1]) + '_vs_' + str(pair[0]) + '_Heatmap.png') == True:
#                document.add_heading(str(pair[1]) + " vs. " + str(pair[0]))
#                document.add_picture('./figures/' + str(pair[1]) + '_vs_' + str(pair[0]) + '_Heatmap.png', height=Inches(6.5), width=Inches(7)) 
#            else:
#                document.add_heading(str(pair[0]) + " vs. " + str(pair[1]))
#                document.add_paragraph('Image not found')
#        
#            
#            document.add_page_break()
            


    #######################################################################################   
    #DOCUMENT COMPLETION

    #Save the document
    docname = 'DataReport_' + data_filename + '.docx'
    document.save(docname)

    end_time = datetime.datetime.now()
    print("Run Time - Document Generation: ", end_time - start_time)
    print("Document generation completed.")

                                   

#####################################################################################
#RUN THE REPORT_GENERATOR FUNCTION
report_generator(df) #default option
#report_generator(df, corrplots = False, binaryplots = False) #no extra plots
                                               
#Additional examples:
#Short printout focusing on only 8 columns
#report_generator(df, [['BOOKING_RLS_DT', 'CONTAINER_EMPTY_AVAIL_DT', 'LOB_SAIL_DT'],['CARRIER', 'SHPMT_UNT_ID', 'TAC'],['DAYS_DIF', 'DT_CNT']])
#
#Short printout with just measure columns
#report_generator(df, None, 'CLEAN SHORT 100000 TEST', ['measure'], corrplots = False, binaryplots = False)  
#
#Just removing the title page                                           
#report_generator(df=df,title='CLEAN SHORT 100000 TEST')
